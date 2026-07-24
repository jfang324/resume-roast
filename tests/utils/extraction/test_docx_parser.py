"""Tests for DocxParser."""

# python-docx ships without complete type stubs in this environment.
# pyright: reportUnknownMemberType=false, reportUnknownArgumentType=false, reportUnknownVariableType=false, reportUnknownParameterType=false, reportGeneralTypeIssues=false

import zipfile
from pathlib import Path

import pytest
from docx import Document

from resume_roast.utils.extraction.document_parser import DocumentParser
from resume_roast.utils.extraction.docx_parser import DocxParser
from resume_roast.utils.extraction.errors import UnreadableDocumentError
from resume_roast.utils.extraction.types import ParsedResume


def _build_sample_document() -> Document:
    """Build an in-memory DOCX resume with the fields the parser tests against."""
    document = Document()
    document.add_heading("Jane Doe", level=1)
    document.add_heading("Experience", level=2)
    document.add_paragraph("Roasted resumes at Acme Corp")
    document.add_heading("Education", level=2)
    document.add_paragraph("BS Computer Science, State University")
    document.core_properties.author = "unit-test"
    return document


@pytest.fixture(scope="module")
def sample_docx(tmp_path_factory: pytest.TempPathFactory) -> Path:
    path = tmp_path_factory.mktemp("extraction") / "sample.docx"
    _build_sample_document().save(path)
    return path


@pytest.fixture(scope="module")
def parsed(sample_docx: Path) -> ParsedResume:
    return DocxParser().parse(sample_docx)


def test_docx_parser_satisfies_document_parser() -> None:
    parser: DocumentParser = DocxParser()
    assert isinstance(parser, DocxParser)


def test_markdown_keeps_inserted_text(parsed: ParsedResume) -> None:
    assert "Jane Doe" in parsed.markdown
    assert "Experience" in parsed.markdown
    assert "Roasted resumes at Acme Corp" in parsed.markdown
    assert "Education" in parsed.markdown


def test_markdown_promotes_headings_to_h1_and_h2(parsed: ParsedResume) -> None:
    assert "# Jane Doe" in parsed.markdown
    assert "## Experience" in parsed.markdown
    assert "## Education" in parsed.markdown


def test_metadata_leaves_pages_empty(parsed: ParsedResume) -> None:
    metadata = parsed.metadata
    assert metadata.page_count == 0
    assert metadata.pages == ()


def test_metadata_extracts_creator_and_records_dates(parsed: ParsedResume) -> None:
    metadata = parsed.metadata
    assert metadata.creator == "unit-test"
    assert metadata.producer is not None
    assert isinstance(metadata.created, str) and "T" in metadata.created
    assert isinstance(metadata.modified, str) and "T" in metadata.modified


def test_metadata_leaves_links_empty(parsed: ParsedResume) -> None:
    assert parsed.metadata.links == ()


def test_blank_core_field_normalizes_to_none(tmp_path: Path) -> None:
    path = tmp_path / "no-author.docx"
    document = _build_sample_document()
    document.core_properties.author = ""
    document.save(path)
    parsed = DocxParser().parse(path)
    assert parsed.metadata.creator is None


def test_missing_file_raises(tmp_path: Path) -> None:
    with pytest.raises(UnreadableDocumentError):
        DocxParser().parse(tmp_path / "missing.docx")


def test_corrupt_file_raises(tmp_path: Path) -> None:
    path = tmp_path / "corrupt.docx"
    path.write_bytes(b"this is not a docx")
    with pytest.raises(UnreadableDocumentError):
        DocxParser().parse(path)


def test_missing_core_xml_part_yields_none_creator(tmp_path: Path) -> None:
    path = tmp_path / "no-core.docx"
    document = _build_sample_document()
    full = tmp_path / "full.docx"
    document.save(full)
    # Re-pack the zip without `docProps/core.xml`; mammoth only needs the body
    # part so this is enough to exercise `_read_part_properties`'s missing-part
    # branch while leaving the rest of the document intact.
    with zipfile.ZipFile(full, "r") as src, zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as dst:
        for name in src.namelist():
            if name == "docProps/core.xml":
                continue
            dst.writestr(name, src.read(name))

    parsed = DocxParser().parse(path)

    assert "Jane Doe" in parsed.markdown
    assert parsed.metadata.creator is None
    assert parsed.metadata.created is None
    assert parsed.metadata.modified is None
    assert parsed.metadata.producer is not None
