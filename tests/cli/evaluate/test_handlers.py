"""Tests for `resume-roast evaluate`."""

# The fixture drives PyMuPDF's partially annotated document-building API.
# python-docx's stub is incomplete in this environment.
# pyright: reportUnknownMemberType=false, reportUnknownArgumentType=false, reportUnknownVariableType=false, reportGeneralTypeIssues=false

import json
import re
from collections.abc import Sequence
from pathlib import Path
from typing import ClassVar

import pymupdf
import pytest
from docx import Document
from typer.testing import CliRunner

from resume_roast.cli.registry import build_subcommand_registry
from resume_roast.integrations.errors import TransientError, TruncatedResponseError
from resume_roast.integrations.types import Completion, Message, Usage
from resume_roast.persistence.credentials.store import CredentialsStore
from resume_roast.persistence.credentials.types import Credentials
from resume_roast.prompts.evaluate.output.schema import CATEGORY_NAMES

app = build_subcommand_registry()
runner = CliRunner()

_MODEL = "nvidia/nemotron-3-super-120b-a12b"


def _report_json() -> str:
    suggestions = [
        {
            "recommendation": "Quantify your impact",
            "examples": [{"quote": "Roasted resumes", "rewrite": "Roasted [X]% more resumes"}],
        }
    ]
    return json.dumps(
        {
            "overall": "It's a roast.",
            "overall_score": 4,
            "categories": {
                name: {"score": 5, "findings": f"{name} needs work.", "suggestions": suggestions}
                for name in CATEGORY_NAMES
            },
            "strengths": ["Concise"],
            "weaknesses": ["No metrics"],
        }
    )


class _FakeClient:
    """Stands in for NvidiaClient; answers prompt() from a queue of texts."""

    texts: ClassVar[list[str]] = []
    """Responses served in order; an empty queue serves the default report."""

    usage: ClassVar[Usage | None] = None
    error: ClassVar[Exception | None] = None
    last: ClassVar["_FakeClient | None"] = None

    def __init__(self, api_key: str, model: str) -> None:
        self.api_key = api_key
        self.model = model
        self.calls: list[list[Message]] = []
        type(self).last = self

    def prompt(
        self,
        messages: Sequence[Message],
        *,
        temperature: float = 0.0,  # noqa: ARG002 — the protocol's signature requires it
    ) -> Completion:
        self.calls.append(list(messages))
        error = type(self).error
        if error is not None:
            raise error
        queue = type(self).texts
        text = queue.pop(0) if queue else _report_json()
        return Completion(text=text, usage=type(self).usage, finish_reason="stop")


@pytest.fixture(autouse=True)
def _isolated_storage_dir(  # pyright: ignore[reportUnusedFunction]
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> Path:
    monkeypatch.setattr("resume_roast.cli.utils.storage_dir", lambda: tmp_path)
    return tmp_path


@pytest.fixture(autouse=True)
def _fake_client(monkeypatch: pytest.MonkeyPatch) -> None:  # pyright: ignore[reportUnusedFunction]
    monkeypatch.setattr("resume_roast.cli.utils.NvidiaClient", _FakeClient)
    monkeypatch.setattr(_FakeClient, "texts", [])
    monkeypatch.setattr(
        _FakeClient,
        "usage",
        Usage(prompt_tokens=2_000, completion_tokens=1_214, total_tokens=3_214),
    )
    monkeypatch.setattr(_FakeClient, "error", None)
    monkeypatch.setattr(_FakeClient, "last", None)


@pytest.fixture
def saved_key(tmp_path: Path) -> None:
    credentials = Credentials(nvidia_api_key="nv-key")  # pragma: allowlist secret
    CredentialsStore(tmp_path).save(credentials)


@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    path = tmp_path / "sample.pdf"
    with pymupdf.open() as doc:
        page = doc.new_page()
        page.insert_text((72, 80), "Jane Doe", fontsize=20)
        page.insert_text((72, 120), "Roasted resumes at Acme Corp", fontsize=11)
        doc.save(path)
    return path


@pytest.mark.usefixtures("saved_key")
def test_evaluate_prints_the_rendered_report(sample_pdf: Path) -> None:
    result = runner.invoke(app, ["evaluate", str(sample_pdf)])

    assert result.exit_code == 0
    assert "[Overall Assessment]\nIt's a roast." in result.output
    assert "Overall: 4/10" in result.output
    for name in CATEGORY_NAMES:
        assert f"[{name} — 5/10]" in result.output
        assert f"{name} needs work." in result.output
    assert "[What's Good]\n- Concise" in result.output
    assert "[What's Bad]\n- No metrics" in result.output
    assert "Suggestions:\n- Quantify your impact" in result.output
    # Diff hunks: the removal line, then the addition line the handler colors.
    assert "  - Roasted resumes" in result.output
    assert "  + Roasted [X]% more resumes" in result.output


@pytest.mark.usefixtures("saved_key")
def test_evaluate_sends_system_and_user_messages(sample_pdf: Path) -> None:
    runner.invoke(app, ["evaluate", str(sample_pdf)])

    client = _FakeClient.last
    assert client is not None
    assert client.api_key == "nv-key"  # pragma: allowlist secret
    assert client.model == _MODEL  # default settings
    system, user = client.calls[0]
    assert system.role == "system"
    assert "## Persona: Recruiter" in system.content
    assert user.role == "user"
    assert "Jane Doe" in user.content


@pytest.mark.usefixtures("saved_key")
def test_evaluate_retries_a_malformed_response(
    sample_pdf: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(_FakeClient, "texts", ["not json at all", _report_json()])

    result = runner.invoke(app, ["evaluate", str(sample_pdf)])

    assert result.exit_code == 0
    assert "Overall: 4/10" in result.output
    client = _FakeClient.last
    assert client is not None
    assert len(client.calls) == 2
    retry_feedback = client.calls[1][-1]
    assert retry_feedback.role == "user"
    assert "not valid JSON" in retry_feedback.content


@pytest.mark.usefixtures("saved_key")
def test_evaluate_reports_a_response_that_never_parses(
    sample_pdf: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(_FakeClient, "texts", ["nope", "still nope"])

    result = runner.invoke(app, ["evaluate", str(sample_pdf)])

    assert result.exit_code == 1
    assert "not valid JSON" in result.output
    assert "Traceback" not in result.output


@pytest.mark.usefixtures("saved_key")
def test_evaluate_reports_truncation_after_retrying(
    sample_pdf: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(
        _FakeClient, "error", TruncatedResponseError("Response hit the completion limit.")
    )

    result = runner.invoke(app, ["evaluate", str(sample_pdf)])

    assert result.exit_code == 1
    assert "completion limit" in result.output
    client = _FakeClient.last
    assert client is not None
    assert len(client.calls) == 2  # the truncation retry ran
    assert "Traceback" not in result.output


@pytest.mark.usefixtures("saved_key")
def test_evaluate_prints_summary_line(sample_pdf: Path) -> None:
    result = runner.invoke(app, ["evaluate", str(sample_pdf)])

    # 2000/1M * $0.09 + 1214/1M * $0.45 = $0.00072...
    assert "2,000 input tokens · 1,214 output tokens · ~$0.0007 · " in result.output
    assert _MODEL not in result.output
    assert re.search(r"· \d+\.\ds", result.output)


@pytest.mark.usefixtures("saved_key")
def test_evaluate_summary_omits_tokens_and_cost_without_usage(
    sample_pdf: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(_FakeClient, "usage", None)

    result = runner.invoke(app, ["evaluate", str(sample_pdf)])

    assert result.exit_code == 0
    assert " input tokens ·" not in result.output
    assert "$" not in result.output
    assert re.search(r"^\d+\.\ds", result.output, re.MULTILINE)


def test_evaluate_requires_an_api_key(sample_pdf: Path) -> None:
    result = runner.invoke(app, ["evaluate", str(sample_pdf)])

    assert result.exit_code == 1
    assert "No NVIDIA API key configured" in result.output
    assert "resume-roast config credentials" in result.output
    assert "Traceback" not in result.output


@pytest.mark.usefixtures("saved_key")
def test_evaluate_reports_client_errors(sample_pdf: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(_FakeClient, "error", TransientError("NVIDIA API is unavailable."))

    result = runner.invoke(app, ["evaluate", str(sample_pdf)])

    assert result.exit_code == 1
    assert "NVIDIA API is unavailable" in result.output
    assert "Traceback" not in result.output


@pytest.mark.usefixtures("saved_key")
def test_evaluate_reports_unreadable_file(tmp_path: Path) -> None:
    path = tmp_path / "corrupt.pdf"
    path.write_bytes(b"this is not a pdf")

    result = runner.invoke(app, ["evaluate", str(path)])

    assert result.exit_code == 1
    assert "Error" in result.output
    assert "Traceback" not in result.output


@pytest.mark.usefixtures("saved_key")
def test_evaluate_reports_missing_file(tmp_path: Path) -> None:
    result = runner.invoke(app, ["evaluate", str(tmp_path / "missing.pdf")])

    assert result.exit_code == 1
    assert "Error" in result.output
    assert "Traceback" not in result.output


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("Jane Doe", level=1)
    document.add_paragraph("Roasted resumes at Acme Corp")
    document.core_properties.author = "unit-test"
    document.save(str(path))
    return path


@pytest.mark.usefixtures("saved_key")
def test_evaluate_accepts_docx_input(sample_docx: Path) -> None:
    result = runner.invoke(app, ["evaluate", str(sample_docx)])

    assert result.exit_code == 0
    assert "Overall: 4/10" in result.output
    assert "[What's Good]\n- Concise" in result.output
    assert "[What's Bad]\n- No metrics" in result.output
    client = _FakeClient.last
    assert client is not None
    _, user = client.calls[0]
    assert user.role == "user"
    assert "Jane Doe" in user.content


@pytest.mark.usefixtures("saved_key")
def test_evaluate_rejects_unsupported_extension(tmp_path: Path) -> None:
    path = tmp_path / "resume.txt"
    path.write_text("not a real resume", encoding="utf-8")

    result = runner.invoke(app, ["evaluate", str(path)])

    assert result.exit_code == 1
    assert "Unsupported file type" in result.output
    assert "Traceback" not in result.output
